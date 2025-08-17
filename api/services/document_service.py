"""
Document service for handling document operations and processing.
"""

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, update, func, delete
from sqlalchemy.orm import selectinload
from typing import List, Optional, Dict, Any
import structlog
import uuid
import os
from datetime import datetime

from models.document import Document, DocumentMetadata, DocumentContent, DocumentProcessingJob
from schemas.document import DocumentCreate, DocumentResponse
from core.minio_client import upload_file_object, delete_file, download_file
from services.llm_service import LLMService
from services.ai_metadata_service import AIMetadataService
import PyPDF2
import docx
import io
import re
from typing import Dict, Any

logger = structlog.get_logger()


class DocumentService:
    """Service for document operations."""
    
    def __init__(self):
        self.llm_service = LLMService()
        self.ai_metadata_service = AIMetadataService()
    
    async def create_document(
        self,
        db: AsyncSession,
        filename: str,
        original_filename: str,
        file_size: int,
        mime_type: str,
        file_extension: str,
        file_path: str
    ) -> Document:
        """Create a new document record."""
        
        try:
            document = Document(
                filename=filename,
                original_filename=original_filename,
                file_path=file_path,
                file_size=file_size,
                mime_type=mime_type,
                file_extension=file_extension,
                status="uploaded"
            )
            
            db.add(document)
            await db.commit()
            await db.refresh(document)
            
            logger.info(f"Document created successfully: {document.id}")
            return document
            
        except Exception as e:
            await db.rollback()
            logger.error(f"Failed to create document: {e}")
            raise
    
    async def get_document(self, db: AsyncSession, document_id: str) -> Optional[Document]:
        """Get a document by ID."""
        
        try:
            query = select(Document).options(
                selectinload(Document.document_metadata),
                selectinload(Document.content)
            ).where(Document.id == document_id)
            
            result = await db.execute(query)
            document = result.scalar_one_or_none()
            
            if document:
                logger.info(f"Document retrieved: {document_id}")
            else:
                logger.warning(f"Document not found: {document_id}")
            
            return document
            
        except Exception as e:
            logger.error(f"Failed to get document {document_id}: {e}")
            raise
    
    async def list_documents(
        self,
        db: AsyncSession,
        skip: int = 0,
        limit: int = 100,
        status: Optional[str] = None,
        agreement_type: Optional[str] = None,
        jurisdiction: Optional[str] = None
    ) -> List[Document]:
        """List documents with optional filtering."""
        
        try:
            query = select(Document).options(
                selectinload(Document.document_metadata)
            )
            
            # Apply filters
            if status:
                query = query.where(Document.status == status)
            if agreement_type:
                query = query.join(DocumentMetadata).where(
                    DocumentMetadata.agreement_type == agreement_type
                )
            if jurisdiction:
                query = query.join(DocumentMetadata).where(
                    DocumentMetadata.jurisdiction == jurisdiction
                )
            
            # Apply pagination
            query = query.offset(skip).limit(limit)
            
            result = await db.execute(query)
            documents = result.scalars().all()
            
            logger.info(f"Retrieved {len(documents)} documents")
            return documents
            
        except Exception as e:
            logger.error(f"Failed to list documents: {e}")
            raise
    
    async def update_document_status(
        self,
        db: AsyncSession,
        document_id: str,
        status: str,
        error_message: Optional[str] = None
    ) -> bool:
        """Update document status."""
        
        try:
            query = update(Document).where(Document.id == document_id).values(
                status=status,
                processing_error=error_message,
                updated_at=datetime.utcnow()
            )
            
            if status == "processing":
                query = query.values(processing_started_at=datetime.utcnow())
            elif status in ["completed", "failed"]:
                query = query.values(processing_completed_at=datetime.utcnow())
            
            result = await db.execute(query)
            await db.commit()
            
            if result.rowcount > 0:
                logger.info(f"Document status updated: {document_id} -> {status}")
                return True
            else:
                logger.warning(f"Document not found for status update: {document_id}")
                return False
                
        except Exception as e:
            await db.rollback()
            logger.error(f"Failed to update document status: {e}")
            raise
    
    async def delete_document(self, db: AsyncSession, document_id: str) -> bool:
        """Soft delete a document."""
        
        try:
            document = await self.get_document(db, document_id)
            if not document:
                return False
            
            # Soft delete
            document.soft_delete()
            await db.commit()
            
            logger.info(f"Document soft deleted: {document_id}")
            return True
            
        except Exception as e:
            await db.rollback()
            logger.error(f"Failed to delete document: {e}")
            raise
    
    async def create_processing_job(
        self,
        db: AsyncSession,
        document_id: str,
        job_type: str,
        priority: int = 0
    ) -> DocumentProcessingJob:
        """Create a document processing job."""
        
        try:
            job = DocumentProcessingJob(
                document_id=document_id,
                job_type=job_type,
                priority=priority,
                status="pending"
            )
            
            db.add(job)
            await db.commit()
            await db.refresh(job)
            
            logger.info(f"Processing job created: {job.id} for document {document_id}")
            return job
            
        except Exception as e:
            await db.rollback()
            logger.error(f"Failed to create processing job: {e}")
            raise
    
    async def extract_text_from_document(
        self,
        db: AsyncSession,
        document_id: str
    ) -> Optional[DocumentContent]:
        """Extract text content from a document with status tracking."""
        
        try:
            # Check if content already exists
            existing_content_query = select(DocumentContent).where(DocumentContent.document_id == document_id)
            result = await db.execute(existing_content_query)
            existing_content = result.scalar_one_or_none()
            
            if existing_content:
                logger.info(f"Content already exists for document: {document_id}")
                return existing_content
            
            # Get document and update status
            document = await self.get_document(db, document_id)
            if not document:
                logger.error(f"Document not found: {document_id}")
                return None
            
            # Update document status to indicate text extraction
            document.status = 'extracting_text'
            await db.commit()
            await db.refresh(document)
            
            # Download file from MinIO
            logger.info(f"Downloading file for text extraction: {document.file_path}")
            file_content = download_file(document.file_path)
            if not file_content:
                logger.error(f"Failed to download file: {document.file_path}")
                # Update status to failed
                document.status = 'failed'
                document.processing_error = "Failed to download file for text extraction"
                await db.commit()
                await db.refresh(document)
                return None
            
            # Extract text based on file type
            logger.info(f"Extracting text from {document.file_extension} file")
            text_content = ""
            if document.file_extension.lower() == "pdf":
                text_content = self._extract_text_from_pdf(file_content)
            elif document.file_extension.lower() in ["docx", "doc"]:
                text_content = self._extract_text_from_docx(file_content)
            else:
                logger.warning(f"Unsupported file type: {document.file_extension}")
                # Update status to failed
                document.status = 'failed'
                document.processing_error = f"Unsupported file type: {document.file_extension}"
                await db.commit()
                await db.refresh(document)
                return None
            
            if not text_content:
                logger.warning(f"No text extracted from document: {document_id}")
                # Update status to failed
                document.status = 'failed'
                document.processing_error = "No text content extracted"
                await db.commit()
                await db.refresh(document)
                return None
            
            # Create document content record
            content = DocumentContent(
                document_id=document_id,
                text_content=text_content,
                word_count=len(text_content.split()),
                character_count=len(text_content),
                extraction_method=f"python_{document.file_extension.lower()}",
                confidence_score=0.8,  # Basic confidence for now
                language_detected="en"  # Default to English
            )
            
            db.add(content)
            await db.commit()
            await db.refresh(content)
            
            # Update document status to indicate text extraction completed
            document.status = 'text_extracted'
            await db.commit()
            await db.refresh(document)
            
            logger.info(f"Text extracted successfully from document: {document_id}")
            return content
            
        except Exception as e:
            await db.rollback()
            logger.error(f"Failed to extract text from document {document_id}: {e}")
            
            # Update document status to failed
            try:
                document = await self.get_document(db, document_id)
                if document:
                    document.status = 'failed'
                    document.processing_error = f"Text extraction failed: {str(e)}"
                    await db.commit()
                    await db.refresh(document)
            except Exception as update_error:
                logger.error(f"Failed to update document status: {update_error}")
            
            raise
    
    def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to extract text from PDF: {e}")
            return ""
    
    def _extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from Word document."""
        try:
            doc = docx.Document(io.BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to extract text from Word document: {e}")
            return ""
    
    async def extract_metadata_from_document(
        self,
        db: AsyncSession,
        document_id: str
    ) -> Optional[DocumentMetadata]:
        """Extract metadata from a document with status tracking."""
        
        try:
            # Check if metadata already exists
            existing_metadata_query = select(DocumentMetadata).where(DocumentMetadata.document_id == document_id)
            result = await db.execute(existing_metadata_query)
            existing_metadata = result.scalar_one_or_none()
            
            if existing_metadata:
                logger.info(f"Metadata already exists for document: {document_id}")
                return existing_metadata
            
            # Get document and update status
            document = await self.get_document(db, document_id)
            if not document:
                logger.error(f"Document not found: {document_id}")
                return None
            
            # Update document status to indicate metadata extraction
            document.status = 'extracting_metadata'
            await db.commit()
            await db.refresh(document)
            
            # Get document content for analysis
            content_query = select(DocumentContent).where(DocumentContent.document_id == document_id)
            result = await db.execute(content_query)
            content = result.scalar_one_or_none()
            
            if not content:
                logger.warning(f"No content found for metadata extraction: {document_id}")
                # Update status to failed
                document.status = 'failed'
                document.processing_error = "No content found for metadata extraction"
                await db.commit()
                await db.refresh(document)
                return None
            
            # Try AI-powered extraction first
            metadata = None
            extraction_method = "rule_based"
            
            if self.ai_metadata_service.client:
                logger.info(f"Attempting AI-powered metadata extraction for {document_id}")
                metadata = self.ai_metadata_service.extract_metadata_with_ai(
                    content.text_content, document.original_filename
                )
                
                if metadata and self.ai_metadata_service.validate_metadata(metadata):
                    metadata = self.ai_metadata_service.enhance_metadata(metadata)
                    extraction_method = "ai_powered"
                    logger.info(f"AI metadata extraction successful for {document_id}")
                else:
                    logger.warning(f"AI metadata extraction failed or invalid for {document_id}, falling back to rule-based")
            
            # Fall back to rule-based extraction if AI fails
            if not metadata:
                logger.info(f"Using rule-based metadata extraction for {document_id}")
                metadata = self._extract_basic_metadata(content.text_content, document.original_filename)
            
            # Filter metadata to only include valid DocumentMetadata fields
            valid_fields = {
                'agreement_type', 'jurisdiction', 'governing_law', 'geography', 
                'industry_sector', 'parties', 'effective_date', 'expiration_date',
                'contract_value', 'currency', 'keywords', 'tags', 'summary',
                'extraction_confidence'
            }
            
            filtered_metadata = {k: v for k, v in metadata.items() if k in valid_fields}
            
            # Create metadata record
            doc_metadata = DocumentMetadata(
                document_id=document_id,
                extraction_method=extraction_method,
                **filtered_metadata
            )
            
            db.add(doc_metadata)
            await db.commit()
            await db.refresh(doc_metadata)
            
            # Update document status to indicate metadata extraction completed
            document.status = 'metadata_extracted'
            await db.commit()
            await db.refresh(document)
            
            logger.info(f"Metadata extracted successfully from document: {document_id}")
            return doc_metadata
            
        except Exception as e:
            await db.rollback()
            logger.error(f"Failed to extract metadata from document {document_id}: {e}")
            
            # Update document status to failed
            try:
                document = await self.get_document(db, document_id)
                if document:
                    document.status = 'failed'
                    document.processing_error = f"Metadata extraction failed: {str(e)}"
                    await db.commit()
                    await db.refresh(document)
            except Exception as update_error:
                logger.error(f"Failed to update document status: {update_error}")
            
            raise
    
    def _extract_basic_metadata(self, text_content: str, filename: str) -> Dict[str, Any]:
        """Extract comprehensive metadata using advanced rule-based patterns."""
        metadata = {}
        text_lower = text_content.lower()
        filename_lower = filename.lower()
        
        # 1. AGREEMENT TYPE EXTRACTION (Enhanced)
        agreement_type_patterns = {
            "NDA": [
                r"non.?disclosure\s+agreement",
                r"nda\s+agreement",
                r"confidentiality\s+agreement",
                r"non.?disclosure\s+and\s+confidentiality"
            ],
            "MSA": [
                r"master\s+services?\s+agreement",
                r"msa\s+agreement",
                r"master\s+agreement"
            ],
            "Franchise Agreement": [
                r"franchise\s+agreement",
                r"franchising\s+agreement",
                r"franchise\s+contract"
            ],
            "Employment Agreement": [
                r"employment\s+agreement",
                r"employment\s+contract",
                r"employment\s+terms"
            ],
            "Tenancy Agreement": [
                r"tenancy\s+agreement",
                r"lease\s+agreement",
                r"rental\s+agreement",
                r"tenancy\s+contract"
            ],
            "Service Agreement": [
                r"service\s+agreement",
                r"consulting\s+agreement",
                r"professional\s+services\s+agreement"
            ],
            "License Agreement": [
                r"license\s+agreement",
                r"licensing\s+agreement",
                r"software\s+license"
            ],
            "Partnership Agreement": [
                r"partnership\s+agreement",
                r"joint\s+venture\s+agreement",
                r"collaboration\s+agreement"
            ],
            "Purchase Agreement": [
                r"purchase\s+agreement",
                r"sales\s+agreement",
                r"acquisition\s+agreement"
            ]
        }
        
        # Check filename first, then content
        for agreement_type, patterns in agreement_type_patterns.items():
            if any(pattern.replace(r"\s+", " ").replace(r"\w+", "").lower() in filename_lower for pattern in patterns):
                metadata["agreement_type"] = agreement_type
                break
            else:
                for pattern in patterns:
                    if re.search(pattern, text_lower):
                        metadata["agreement_type"] = agreement_type
                        break
                if "agreement_type" in metadata:
                    break
        
        # 2. JURISDICTION & GOVERNING LAW EXTRACTION (Enhanced)
        jurisdiction_patterns = [
            # Specific countries and states - more targeted
            r"governed\s+by\s+the\s+laws?\s+of\s+the\s+([A-Za-z]+(?:\s+[A-Za-z]+)*?)(?:\s|,|\.|$)",
            r"governed\s+by\s+the\s+laws?\s+of\s+([A-Za-z]+(?:\s+[A-Za-z]+)*?)(?:\s|,|\.|$)",
            r"jurisdiction\s+of\s+([A-Za-z]+(?:\s+[A-Za-z]+)*?)(?:\s|,|\.|$)",
            r"subject\s+to\s+([A-Za-z]+(?:\s+[A-Za-z]+)*?)\s+law",
            r"([A-Za-z]+(?:\s+[A-Za-z]+)*?)\s+law\s+shall\s+govern",
            r"governed\s+by\s+([A-Za-z]+(?:\s+[A-Za-z]+)*?)\s+law",
            r"([A-Za-z]+(?:\s+[A-Za-z]+)*?)\s+courts?\s+shall\s+have\s+exclusive\s+jurisdiction",
            r"exclusive\s+jurisdiction\s+of\s+([A-Za-z]+(?:\s+[A-Za-z]+)*?)(?:\s|,|\.|$)",
            r"venue\s+shall\s+be\s+([A-Za-z]+(?:\s+[A-Za-z]+)*?)(?:\s|,|\.|$)",
            r"([A-Za-z]+(?:\s+[A-Za-z]+)*?)\s+venue",
            r"([A-Za-z]+(?:\s+[A-Za-z]+)*?)\s+governing\s+law",
            # Additional patterns for specific legal language
            r"laws\s+of\s+the\s+([A-Za-z]+(?:\s+[A-Za-z]+)*?)(?:\s|,|\.|$)",
            r"courts\s+of\s+([A-Za-z]+(?:\s+[A-Za-z]+)*?)(?:\s|,|\.|$)"
        ]
        
        # Common jurisdictions mapping
        jurisdiction_mapping = {
            "uae": "UAE", "united arab emirates": "UAE",
            "uk": "UK", "united kingdom": "UK", "england": "UK", "wales": "UK",
            "usa": "USA", "united states": "USA", "us": "USA",
            "delaware": "Delaware, USA", "california": "California, USA", "new york": "New York, USA",
            "singapore": "Singapore", "hong kong": "Hong Kong",
            "germany": "Germany", "france": "France", "netherlands": "Netherlands",
            "australia": "Australia", "canada": "Canada", "japan": "Japan",
            # Additional mappings for common legal jurisdictions
            "state of delaware": "Delaware, USA", "delaware state": "Delaware, USA",
            "state of california": "California, USA", "california state": "California, USA",
            "state of new york": "New York, USA", "new york state": "New York, USA"
        }
        
        for pattern in jurisdiction_patterns:
            match = re.search(pattern, text_lower)
            if match:
                jurisdiction = match.group(1).strip().lower()
                # Map to standard format
                for key, value in jurisdiction_mapping.items():
                    if key in jurisdiction:
                        metadata["jurisdiction"] = value
                        metadata["governing_law"] = value
                        break
                else:
                    metadata["jurisdiction"] = jurisdiction.title()
                    metadata["governing_law"] = jurisdiction.title()
                break
        
        # 3. GEOGRAPHY EXTRACTION (Enhanced)
        geography_patterns = [
            r"(\w+(?:\s+\w+)*)\s+region",
            r"(\w+(?:\s+\w+)*)\s+territory",
            r"(\w+(?:\s+\w+)*)\s+state",
            r"(\w+(?:\s+\w+)*)\s+country",
            r"(\w+(?:\s+\w+)*)\s+area",
            r"(\w+(?:\s+\w+)*)\s+zone",
            r"(\w+(?:\s+\w+)*)\s+district",
            r"(\w+(?:\s+\w+)*)\s+province",
            r"(\w+(?:\s+\w+)*)\s+county"
        ]
        
        # Geography mapping for standardization
        geography_mapping = {
            "middle east": "Middle East", "gulf region": "Gulf Region", "gcc": "Gulf Cooperation Council",
            "europe": "Europe", "european union": "European Union", "eu": "European Union",
            "asia": "Asia", "asia pacific": "Asia Pacific", "apac": "Asia Pacific",
            "north america": "North America", "south america": "South America",
            "africa": "Africa", "australia": "Australia", "oceania": "Oceania"
        }
        
        for pattern in geography_patterns:
            match = re.search(pattern, text_lower)
            if match:
                geography = match.group(1).strip().lower()
                # Filter out section headers and common legal words
                if (geography not in ["governing", "law", "and", "jurisdiction", "this", "agreement", "shall", "be", "governed", "by", "construed", "accordance", "with", "laws", "of", "the", "state", "united", "states", "america", "disputes", "arising", "out", "relating", "subject", "exclusive", "courts", "parties", "concerning", "subject", "matter", "hereof", "supersedes", "prior", "contemporaneous", "agreements", "understandings", "whether", "written", "oral", "relating", "such", "subject", "matter"] and
                    len(geography) > 3):
                    # Map to standard format
                    for key, value in geography_mapping.items():
                        if key in geography:
                            metadata["geography"] = value
                            break
                    else:
                        metadata["geography"] = geography.title()
                    break
        
        # 4. INDUSTRY SECTOR EXTRACTION (Enhanced)
        industry_patterns = {
            "Oil & Gas": [
                r"oil\s+and\s+gas", r"petroleum", r"hydrocarbon", r"drilling", r"exploration",
                r"upstream", r"downstream", r"midstream", r"refinery", r"petrochemical"
            ],
            "Healthcare": [
                r"healthcare", r"medical", r"pharmaceutical", r"biotech", r"clinical",
                r"hospital", r"diagnostic", r"therapeutic", r"medical device"
            ],
            "Technology": [
                r"technology", r"software", r"hardware", r"it\s+services", r"digital",
                r"cybersecurity", r"artificial intelligence", r"machine learning", r"cloud"
            ],
            "Finance": [
                r"finance", r"banking", r"investment", r"insurance", r"asset management",
                r"private equity", r"venture capital", r"fintech", r"wealth management"
            ],
            "Real Estate": [
                r"real estate", r"property", r"construction", r"development", r"leasing",
                r"commercial property", r"residential", r"infrastructure"
            ],
            "Manufacturing": [
                r"manufacturing", r"industrial", r"production", r"factory", r"supply chain",
                r"logistics", r"automotive", r"aerospace", r"chemical"
            ],
            "Retail": [
                r"retail", r"e-commerce", r"consumer goods", r"fashion", r"apparel",
                r"food and beverage", r"hospitality", r"tourism"
            ],
            "Energy": [
                r"energy", r"renewable", r"solar", r"wind", r"nuclear", r"electricity",
                r"power generation", r"utilities"
            ]
        }
        
        for industry, patterns in industry_patterns.items():
            for pattern in patterns:
                if re.search(pattern, text_lower):
                    metadata["industry_sector"] = industry
                    break
            if "industry_sector" in metadata:
                break
        
        # 5. PARTIES EXTRACTION (Enhanced)
        parties_patterns = [
            r"between\s+([^:]+?)\s+and\s+([^:]+?)(?:\s|,|\.|$)",
            r"parties\s+are\s+([^:]+?)\s+and\s+([^:]+?)(?:\s|,|\.|$)",
            r"([^:]+?)\s+\(hereinafter\s+referred\s+to\s+as\s+[^)]+\)",
            r"([^:]+?)\s+\(the\s+[^)]+\)",
            r"landlord[:\s]+([^,\n]+)",
            r"tenant[:\s]+([^,\n]+)",
            r"buyer[:\s]+([^,\n]+)",
            r"seller[:\s]+([^,\n]+)",
            r"licensor[:\s]+([^,\n]+)",
            r"licensee[:\s]+([^,\n]+)"
        ]
        
        parties = []
        for pattern in parties_patterns:
            matches = re.findall(pattern, text_content, re.IGNORECASE)
            for match in matches:
                if isinstance(match, tuple):
                    parties.extend([party.strip() for party in match if party.strip()])
                else:
                    parties.append(match.strip())
        
        # Clean and standardize party names
        cleaned_parties = []
        for party in parties:
            # Remove common legal terms and clean up
            party_clean = re.sub(r'\([^)]*\)', '', party)  # Remove parentheses content
            party_clean = re.sub(r'hereinafter\s+referred\s+to\s+as\s+\w+', '', party_clean)
            party_clean = re.sub(r'the\s+', '', party_clean)
            party_clean = party_clean.strip()
            if party_clean and len(party_clean) > 2:
                cleaned_parties.append(party_clean)
        
        if cleaned_parties:
            metadata["parties"] = list(set(cleaned_parties))  # Remove duplicates
        
        # 6. DATE EXTRACTION (Enhanced)
        date_patterns = [
            # Effective dates
            r"effective\s+date[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
            r"commencement\s+date[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
            r"start\s+date[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
            r"execution\s+date[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
            r"signing\s+date[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
            # Expiration dates
            r"expiration\s+date[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
            r"end\s+date[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
            r"termination\s+date[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
            # Date ranges
            r"from\s+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\s+to\s+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
            # Specific date formats
            r"(\d{1,2}\s+(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{4})",
            r"(\d{4}[-/]\d{1,2}[-/]\d{1,2})"
        ]
        
        for pattern in date_patterns:
            matches = re.findall(pattern, text_lower)
            for match in matches:
                if isinstance(match, tuple):
                    # Handle date ranges
                    try:
                        from datetime import datetime
                        date1 = self._parse_date(match[0])
                        date2 = self._parse_date(match[1])
                        if date1:
                            metadata["effective_date"] = date1
                        if date2:
                            metadata["expiration_date"] = date2
                        break
                    except:
                        pass
                else:
                    try:
                        date_obj = self._parse_date(match)
                        if date_obj:
                            if "effective_date" not in metadata:
                                metadata["effective_date"] = date_obj
                            elif "expiration_date" not in metadata:
                                metadata["expiration_date"] = date_obj
                            break
                    except:
                        pass
        
        # 7. CONTRACT VALUE & CURRENCY EXTRACTION (Enhanced)
        currency_patterns = [
            # Dollar amounts
            r"(\$[\d,]+(?:\.\d{2})?)",
            r"(\d+(?:,\d{3})*(?:\.\d{2})?\s*(?:USD|US\s*Dollars?))",
            # Euro amounts
            r"(\d+(?:,\d{3})*(?:\.\d{2})?\s*(?:EUR|Euros?))",
            # Pound amounts
            r"(\d+(?:,\d{3})*(?:\.\d{2})?\s*(?:GBP|Pounds?|Sterling))",
            # AED amounts
            r"(\d+(?:,\d{3})*(?:\.\d{2})?\s*(?:AED|Dirhams?))",
            # Generic amounts with currency
            r"(\d+(?:,\d{3})*(?:\.\d{2})?\s*(?:dollars?|euros?|pounds?|dirhams?))",
            # Written amounts
            r"(?:amount|value|consideration)[:\s]+([^,\n]+)"
        ]
        
        for pattern in currency_patterns:
            match = re.search(pattern, text_content, re.IGNORECASE)
            if match:
                amount_str = match.group(1)
                # Determine currency
                if "$" in amount_str or "USD" in amount_str.upper() or "dollars" in amount_str.lower():
                    metadata["currency"] = "USD"
                elif "EUR" in amount_str.upper() or "euros" in amount_str.lower():
                    metadata["currency"] = "EUR"
                elif "GBP" in amount_str.upper() or "pounds" in amount_str.lower() or "sterling" in amount_str.lower():
                    metadata["currency"] = "GBP"
                elif "AED" in amount_str.upper() or "dirhams" in amount_str.lower():
                    metadata["currency"] = "AED"
                else:
                    metadata["currency"] = "USD"  # Default
                
                # Extract numeric value
                amount_clean = re.sub(r'[^\d.]', '', amount_str)
                try:
                    metadata["contract_value"] = float(amount_clean)
                except:
                    pass
                break
        
        # 8. KEYWORDS EXTRACTION (Enhanced)
        legal_keywords = [
            "confidentiality", "termination", "liability", "indemnification",
            "force majeure", "governing law", "dispute resolution", "breach",
            "remedies", "waiver", "severability", "entire agreement",
            "non-compete", "non-solicitation", "intellectual property",
            "data protection", "privacy", "compliance", "regulatory",
            "audit", "inspection", "default", "cure period",
            "assignment", "amendment", "notice", "representation",
            "warranty", "covenant", "condition precedent", "material adverse effect"
        ]
        
        found_keywords = []
        for keyword in legal_keywords:
            if keyword in text_lower:
                found_keywords.append(keyword)
        
        if found_keywords:
            metadata["keywords"] = found_keywords
        
        # 9. TAGS EXTRACTION (Enhanced)
        tags = []
        if "agreement_type" in metadata:
            tags.append(metadata["agreement_type"])
        if "industry_sector" in metadata:
            tags.append(metadata["industry_sector"])
        if "jurisdiction" in metadata:
            tags.append(metadata["jurisdiction"])
        if "geography" in metadata:
            tags.append(metadata["geography"])
        
        if tags:
            metadata["tags"] = tags
        
        # 10. CONFIDENCE SCORE CALCULATION (Enhanced)
        confidence_score = 0.0
        confidence_factors = {
            "agreement_type": 0.2,
            "jurisdiction": 0.15,
            "geography": 0.1,
            "industry_sector": 0.15,
            "parties": 0.1,
            "effective_date": 0.1,
            "expiration_date": 0.05,
            "contract_value": 0.05,
            "currency": 0.05,
            "keywords": 0.05
        }
        
        for field, weight in confidence_factors.items():
            if field in metadata and metadata[field]:
                confidence_score += weight
        
        metadata["extraction_confidence"] = min(confidence_score, 1.0)
        
        return metadata
    
    def _parse_date(self, date_str: str) -> Optional[datetime]:
        """Parse various date formats and return datetime object."""
        try:
            from datetime import datetime
            
            # Clean the date string
            date_str = date_str.strip()
            
            # Handle different date formats
            date_formats = [
                "%m/%d/%Y", "%d/%m/%Y", "%Y/%m/%d",
                "%m-%d-%Y", "%d-%m-%Y", "%Y-%m-%d",
                "%d %b %Y", "%d %B %Y", "%b %d %Y", "%B %d %Y"
            ]
            
            # Try different formats
            for fmt in date_formats:
                try:
                    return datetime.strptime(date_str, fmt)
                except ValueError:
                    continue
            
            # Handle 2-digit years
            if re.match(r'\d{1,2}[/-]\d{1,2}[/-]\d{2}$', date_str):
                if "/" in date_str:
                    parts = date_str.split("/")
                else:
                    parts = date_str.split("-")
                
                if len(parts) == 3:
                    # Assume MM/DD/YY format
                    month, day, year = parts
                    if len(year) == 2:
                        year = "20" + year
                    return datetime.strptime(f"{month}/{day}/{year}", "%m/%d/%Y")
            
            return None
            
        except Exception:
            return None
    
    async def update_processing_job(
        self,
        db: AsyncSession,
        job_id: str,
        status: str,
        result_data: Optional[Dict[str, Any]] = None,
        error_message: Optional[str] = None
    ) -> bool:
        """Update processing job status."""
        
        try:
            job = await db.get(DocumentProcessingJob, job_id)
            if not job:
                return False
            
            if status == "processing":
                job.start()
            elif status == "completed":
                job.complete(result_data)
            elif status == "failed":
                job.fail(error_message)
            elif status == "retry":
                job.retry()
            
            await db.commit()
            
            logger.info(f"Processing job updated: {job_id} -> {status}")
            return True
            
        except Exception as e:
            await db.rollback()
            logger.error(f"Failed to update processing job: {e}")
            raise
    
    async def mark_document_failed(
        self,
        db: AsyncSession,
        document_id: str,
        error_message: str
    ) -> bool:
        """Mark a document as failed with error message."""
        
        try:
            # Update document status
            update_query = update(Document).where(Document.id == document_id).values(
                status="failed",
                processing_error=error_message,
                processing_completed_at=datetime.utcnow(),
                updated_at=datetime.utcnow()
            )
            
            await db.execute(update_query)
            await db.commit()
            
            logger.error(f"Document marked as failed: {document_id} - {error_message}")
            return True
            
        except Exception as e:
            await db.rollback()
            logger.error(f"Failed to mark document as failed: {e}")
            raise
    
    async def reset_document_status(
        self,
        db: AsyncSession,
        document_id: str,
        status: str = "pending"
    ) -> bool:
        """Reset document status for reprocessing."""
        
        try:
            # Update document status
            update_query = update(Document).where(Document.id == document_id).values(
                status=status,
                processing_error=None,
                processing_started_at=None,
                processing_completed_at=None,
                updated_at=datetime.utcnow()
            )
            
            await db.execute(update_query)
            await db.commit()
            
            logger.info(f"Document status reset: {document_id} -> {status}")
            return True
            
        except Exception as e:
            await db.rollback()
            logger.error(f"Failed to reset document status: {e}")
            raise
    
    async def extract_document_metadata(
        self,
        db: AsyncSession,
        document_id: str,
        text_content: str
    ) -> DocumentMetadata:
        """Extract metadata from document content using LLM service."""
        
        try:
            # Analyze content using LLM service
            analysis = await self.llm_service.analyze_document_content(text_content)
            
            # Create metadata record
            metadata = DocumentMetadata(
                document_id=document_id,
                agreement_type=analysis.get("agreement_type"),
                jurisdiction=analysis.get("jurisdiction"),
                governing_law=analysis.get("jurisdiction"),  # Same as jurisdiction for now
                geography=analysis.get("geography"),
                industry_sector=analysis.get("industry_sector"),
                parties=analysis.get("parties"),
                effective_date=self._parse_date(analysis.get("effective_date")),
                expiration_date=self._parse_date(analysis.get("expiration_date")),
                contract_value=analysis.get("contract_value"),
                currency=analysis.get("currency"),
                keywords=analysis.get("keywords"),
                tags=[],
                extraction_confidence=analysis.get("confidence_score"),
                extraction_method="llm_mock"
            )
            
            db.add(metadata)
            await db.commit()
            await db.refresh(metadata)
            
            logger.info(f"Document metadata extracted: {document_id}")
            return metadata
            
        except Exception as e:
            await db.rollback()
            logger.error(f"Failed to extract metadata: {e}")
            raise
    
    async def store_document_content(
        self,
        db: AsyncSession,
        document_id: str,
        text_content: str,
        extraction_method: str = "manual"
    ) -> DocumentContent:
        """Store extracted document content."""
        
        try:
            # Calculate text statistics
            word_count = len(text_content.split())
            character_count = len(text_content)
            
            # Create content record
            content = DocumentContent(
                document_id=document_id,
                text_content=text_content,
                word_count=word_count,
                character_count=character_count,
                sections=[],
                paragraphs=text_content.split('\n\n'),
                tables=[],
                extraction_method=extraction_method,
                extraction_timestamp=datetime.utcnow(),
                confidence_score=0.9,
                language_detected="en"  # Default to English
            )
            
            db.add(content)
            await db.commit()
            await db.refresh(content)
            
            logger.info(f"Document content stored: {document_id}")
            return content
            
        except Exception as e:
            await db.rollback()
            logger.error(f"Failed to store content: {e}")
            raise
    
    def _parse_date(self, date_str: Optional[str]) -> Optional[datetime]:
        """Parse date string to datetime object."""
        
        if not date_str:
            return None
        
        try:
            # Try different date formats
            date_formats = [
                "%m/%d/%Y",
                "%m/%d/%y",
                "%Y-%m-%d",
                "%d/%m/%Y",
                "%d/%m/%y"
            ]
            
            for fmt in date_formats:
                try:
                    return datetime.strptime(date_str, fmt)
                except ValueError:
                    continue
            
            logger.warning(f"Could not parse date: {date_str}")
            return None
            
        except Exception as e:
            logger.error(f"Date parsing error: {e}")
            return None
    
    async def get_document_statistics(self, db: AsyncSession) -> Dict[str, Any]:
        """Get document statistics."""
        
        try:
            # Total documents
            total_query = select(func.count(Document.id))
            total_result = await db.execute(total_query)
            total_documents = total_result.scalar()
            
            # Documents by status
            status_query = select(
                Document.status,
                func.count(Document.id).label("count")
            ).group_by(Document.status)
            
            status_result = await db.execute(status_query)
            status_counts = {row.status: row.count for row in status_result.all()}
            
            # File size statistics
            size_query = select(
                func.avg(Document.file_size).label("avg_size"),
                func.min(Document.file_size).label("min_size"),
                func.max(Document.file_size).label("max_size")
            )
            
            size_result = await db.execute(size_query)
            size_stats = size_result.first()
            
            return {
                "total_documents": total_documents,
                "status_breakdown": status_counts,
                "file_size_stats": {
                    "average_size_bytes": float(size_stats.avg_size) if size_stats.avg_size else 0,
                    "min_size_bytes": size_stats.min_size or 0,
                    "max_size_bytes": size_stats.max_size or 0
                }
            }
            
        except Exception as e:
            logger.error(f"Failed to get document statistics: {e}")
            raise
    
    async def cleanup_failed_documents(self, db: AsyncSession) -> int:
        """Clean up documents that failed processing multiple times."""
        
        try:
            # Find documents with failed processing jobs
            failed_jobs_query = select(Document.id).join(DocumentProcessingJob).where(
                DocumentProcessingJob.status == "failed",
                DocumentProcessingJob.retry_count >= DocumentProcessingJob.max_retries
            )
            
            failed_result = await db.execute(failed_jobs_query)
            failed_document_ids = [row[0] for row in failed_result.all()]
            
            if not failed_document_ids:
                return 0
            
            # Update status to failed
            update_query = update(Document).where(
                Document.id.in_(failed_document_ids)
            ).values(
                status="failed",
                processing_error="Max retries exceeded"
            )
            
            await db.execute(update_query)
            await db.commit()
            
            logger.info(f"Cleaned up {len(failed_document_ids)} failed documents")
            return len(failed_document_ids)
            
        except Exception as e:
            await db.rollback()
            logger.error(f"Failed to cleanup failed documents: {e}")
            raise 

    async def reprocess_document(
        self,
        db: AsyncSession,
        document_id: str
    ) -> Dict[str, Any]:
        """
        Reprocess a document by clearing existing metadata and content, then re-extracting everything.
        
        Args:
            db: Database session
            document_id: ID of the document to reprocess
            
        Returns:
            Dictionary with reprocessing status and results
        """
        try:
            # Get the document
            document = await self.get_document(db, document_id)
            if not document:
                raise Exception(f"Document not found: {document_id}")
            
            # Check if document is currently being processed
            if document.status == "processing":
                raise Exception("Cannot reprocess document while it is currently being processed")
            
            # Update document status to indicate reprocessing
            document.status = 'processing'
            document.processing_started_at = datetime.utcnow()
            document.processing_error = None
            await db.commit()
            await db.refresh(document)
            
            # Delete existing metadata
            metadata_delete_query = delete(DocumentMetadata).where(DocumentMetadata.document_id == document_id)
            await db.execute(metadata_delete_query)
            
            # Delete existing content
            content_delete_query = delete(DocumentContent).where(DocumentContent.document_id == document_id)
            await db.execute(content_delete_query)
            
            # Delete existing processing jobs
            jobs_delete_query = delete(DocumentProcessingJob).where(DocumentProcessingJob.document_id == document_id)
            await db.execute(jobs_delete_query)
            
            await db.commit()
            
            # Re-extract text content
            logger.info(f"Re-extracting text from document: {document_id}")
            content = await self.extract_text_from_document(db, document_id)
            if not content:
                raise Exception("Failed to re-extract text from document")
            
            # Re-extract metadata
            logger.info(f"Re-extracting metadata from document: {document_id}")
            metadata = await self.extract_metadata_from_document(db, document_id)
            if not metadata:
                raise Exception("Failed to re-extract metadata from document")
            
            # Update document status to completed
            document.status = 'completed'
            document.processing_completed_at = datetime.utcnow()
            await db.commit()
            await db.refresh(document)
            
            logger.info(f"Document reprocessing completed successfully: {document_id}")
            
            return {
                "message": "Document reprocessed successfully",
                "document_id": document_id,
                "filename": document.original_filename,
                "status": "completed",
                "content_id": str(content.id),
                "metadata_id": str(metadata.id)
            }
            
        except Exception as e:
            error_msg = f"Document reprocessing failed: {str(e)}"
            logger.error(error_msg)
            
            # Update document status to failed
            if db and 'document' in locals():
                try:
                    document.status = 'failed'
                    document.processing_error = error_msg
                    await db.commit()
                    await db.refresh(document)
                except Exception as update_error:
                    logger.error(f"Failed to update document status: {update_error}")
            
            raise Exception(error_msg) 